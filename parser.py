"""
parser.py — Resume parsing module for ATS Scanner.
Handles PDF and DOCX resume extraction with structured data output.
"""

import re
import io
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

# PDF parsing
try:
    import pdfplumber
    PDF_BACKEND = "pdfplumber"
except ImportError:
    pdfplumber = None
    PDF_BACKEND = None

# DOCX parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ParsedResume:
    """Structured representation of a parsed resume."""
    raw_text: str = ""
    name: str = ""
    email: str = ""
    phone: str = ""
    linkedin: str = ""
    github: str = ""
    skills: list[str] = field(default_factory=list)
    education: list[dict] = field(default_factory=list)
    experience: list[dict] = field(default_factory=list)
    projects: list[dict] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    summary: str = ""
    total_years_experience: float = 0.0


class ResumeParser:
    """
    Production-grade resume parser supporting PDF and DOCX formats.
    Extracts structured data using regex patterns and heuristics.
    """

    # Common section headers
    SECTION_HEADERS = {
        "experience": r"(?i)(work\s+)?experience|employment(\s+history)?|professional\s+background|career\s+history",
        "education": r"(?i)education|academic(\s+background)?|qualifications",
        "skills": r"(?i)skills|technical\s+skills|core\s+competencies|technologies|expertise",
        "projects": r"(?i)projects|personal\s+projects|key\s+projects|portfolio",
        "certifications": r"(?i)certif(ications?|ied)|licenses?|credentials",
        "summary": r"(?i)(professional\s+)?summary|objective|profile|about\s+me",
    }

    # Weak bullet point starters to flag later
    WEAK_BULLET_PATTERNS = [
        r"^responsible\s+for",
        r"^helped\s+(to\s+)?",
        r"^worked\s+(on|with)",
        r"^assisted\s+(in|with)",
        r"^duties\s+included",
        r"^involved\s+in",
        r"^participated\s+in",
        r"^tasked\s+with",
    ]

    def __init__(self):
        self._weak_bullet_re = [re.compile(p, re.I) for p in self.WEAK_BULLET_PATTERNS]

    # ─────────────────────────────────────────────
    # Public API
    # ─────────────────────────────────────────────

    def parse(self, source: str | Path | bytes, file_type: Optional[str] = None) -> ParsedResume:
        """
        Parse a resume from file path, bytes, or raw text.

        Args:
            source: File path, raw bytes, or plain text string.
            file_type: 'pdf', 'docx', or 'text'. Auto-detected if None.

        Returns:
            ParsedResume dataclass with all extracted fields.
        """
        if isinstance(source, (str, Path)) and Path(source).exists():
            file_type = file_type or Path(source).suffix.lstrip(".").lower()
            raw_text = self._extract_from_file(Path(source), file_type)
        elif isinstance(source, bytes):
            if file_type is None:
                raise ValueError("file_type must be specified when passing bytes.")
            raw_text = self._extract_from_bytes(source, file_type)
        else:
            raw_text = str(source)

        return self._parse_text(raw_text)

    def parse_text(self, text: str) -> ParsedResume:
        """Parse a resume directly from plain text."""
        return self._parse_text(text)

    # ─────────────────────────────────────────────
    # Extraction backends
    # ─────────────────────────────────────────────

    def _extract_from_file(self, path: Path, file_type: str) -> str:
        if file_type == "pdf":
            return self._extract_pdf(path)
        elif file_type in ("docx", "doc"):
            return self._extract_docx(path)
        else:
            return path.read_text(encoding="utf-8", errors="ignore")

    def _extract_from_bytes(self, data: bytes, file_type: str) -> str:
        if file_type == "pdf":
            return self._extract_pdf(io.BytesIO(data))
        elif file_type in ("docx", "doc"):
            return self._extract_docx(io.BytesIO(data))
        return data.decode("utf-8", errors="ignore")

    def _extract_pdf(self, source) -> str:
        """Extract text from PDF using pdfplumber."""
        if pdfplumber is None:
            raise ImportError("pdfplumber is required for PDF parsing. Run: pip install pdfplumber")
        pages = []
        with pdfplumber.open(source) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    pages.append(text)
        return "\n".join(pages)

    def _extract_docx(self, source) -> str:
        """Extract text from DOCX using python-docx."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing. Run: pip install python-docx")
        doc = DocxDocument(source)
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    # ─────────────────────────────────────────────
    # Core parsing logic
    # ─────────────────────────────────────────────

    def _parse_text(self, text: str) -> ParsedResume:
        resume = ParsedResume(raw_text=text)
        resume.name = self._extract_name(text)
        resume.email = self._extract_email(text)
        resume.phone = self._extract_phone(text)
        resume.linkedin = self._extract_linkedin(text)
        resume.github = self._extract_github(text)

        sections = self._split_into_sections(text)

        resume.summary = self._extract_summary(sections)
        resume.skills = self._extract_skills(sections, text)
        resume.education = self._extract_education(sections)
        resume.experience = self._extract_experience(sections)
        resume.projects = self._extract_projects(sections)
        resume.certifications = self._extract_certifications(sections)
        resume.total_years_experience = self._estimate_years_experience(resume.experience)

        return resume

    # ─────────────────────────────────────────────
    # Field extractors
    # ─────────────────────────────────────────────

    def _extract_name(self, text: str) -> str:
        """Heuristic: name is usually on the first non-empty line."""
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            # Skip lines that look like section headers or contact info
            if re.search(r"@|http|linkedin|github|\d{3}[-.\s]\d{3}", line, re.I):
                continue
            if len(line.split()) <= 5 and re.match(r"^[A-Za-z\s\-\.\']+$", line):
                return line.title()
        return ""

    def _extract_email(self, text: str) -> str:
        match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
        return match.group(0) if match else ""

    def _extract_phone(self, text: str) -> str:
        match = re.search(
            r"(\+?\d{1,3}[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}", text
        )
        return match.group(0).strip() if match else ""

    def _extract_linkedin(self, text: str) -> str:
        match = re.search(r"linkedin\.com/in/([A-Za-z0-9\-_]+)", text, re.I)
        return f"linkedin.com/in/{match.group(1)}" if match else ""

    def _extract_github(self, text: str) -> str:
        match = re.search(r"github\.com/([A-Za-z0-9\-_]+)", text, re.I)
        return f"github.com/{match.group(1)}" if match else ""

    def _split_into_sections(self, text: str) -> dict[str, str]:
        """Split resume text into named sections."""
        lines = text.splitlines()
        sections: dict[str, list[str]] = {"header": []}
        current_section = "header"

        for line in lines:
            stripped = line.strip()
            matched_section = None
            for section_name, pattern in self.SECTION_HEADERS.items():
                # Strip inline (?i) flags from pattern and apply re.I at call site
                clean_pattern = re.sub(r"\(\?i\)", "", pattern)
                if re.match(rf"^{clean_pattern}\s*:?\s*$", stripped, re.I):
                    matched_section = section_name
                    break

            if matched_section:
                current_section = matched_section
                sections.setdefault(current_section, [])
            else:
                sections.setdefault(current_section, [])
                sections[current_section].append(line)

        return {k: "\n".join(v) for k, v in sections.items()}

    def _extract_summary(self, sections: dict) -> str:
        return sections.get("summary", "").strip()

    def _extract_skills(self, sections: dict, full_text: str) -> list[str]:
        """Extract skills from skills section and infer from text."""
        skill_text = sections.get("skills", "")
        skills: set[str] = set()

        # Parse comma/pipe/bullet-separated skills
        raw = re.sub(r"[•·▪▸►\-\*]", ",", skill_text)
        for token in re.split(r"[,|\n/]", raw):
            token = token.strip().strip(":")
            if token and 1 < len(token) < 60:
                skills.add(token)

        return sorted(skills)

    def _extract_education(self, sections: dict) -> list[dict]:
        """Extract education entries."""
        text = sections.get("education", "")
        entries = []
        degree_pattern = re.compile(
            r"(B\.?S\.?|B\.?E\.?|B\.?Tech|B\.?A\.?|M\.?S\.?|M\.?Tech|M\.?A\.?|Ph\.?D\.?|MBA|Bachelor|Master|Doctor)[^\n]*",
            re.I,
        )
        year_pattern = re.compile(r"(19|20)\d{2}")

        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            degree_match = degree_pattern.search(line)
            years = year_pattern.findall(line)
            if degree_match or years:
                entries.append({
                    "raw": line,
                    "degree": degree_match.group(0) if degree_match else "",
                    "years": years,
                })

        return entries

    def _extract_experience(self, sections: dict) -> list[dict]:
        """Extract experience blocks with bullet point analysis."""
        text = sections.get("experience", "")
        if not text.strip():
            return []

        entries = []
        year_re = re.compile(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\s*(19|20)\d{2}", re.I)
        bullet_re = re.compile(r"^[•·▪▸►\-\*]\s*(.+)")

        blocks = re.split(r"\n{2,}", text)
        for block in blocks:
            lines = [l.strip() for l in block.splitlines() if l.strip()]
            if not lines:
                continue

            bullets = []
            weak_bullets = []
            for line in lines:
                bm = bullet_re.match(line)
                if bm:
                    bullet_text = bm.group(1)
                    bullets.append(bullet_text)
                    if any(p.match(bullet_text) for p in self._weak_bullet_re):
                        weak_bullets.append(bullet_text)

            years = year_re.findall(block)
            entries.append({
                "raw": block,
                "title_line": lines[0] if lines else "",
                "bullets": bullets,
                "weak_bullets": weak_bullets,
                "years_mentioned": [f"{m[0]} {m[1]}" for m in years] if years else [],
            })

        return entries

    def _extract_projects(self, sections: dict) -> list[dict]:
        """Extract project entries."""
        text = sections.get("projects", "")
        if not text.strip():
            return []

        projects = []
        blocks = re.split(r"\n{2,}", text)
        for block in blocks:
            lines = [l.strip() for l in block.splitlines() if l.strip()]
            if lines:
                projects.append({
                    "raw": block,
                    "title": lines[0],
                    "description": " ".join(lines[1:]),
                })
        return projects

    def _extract_certifications(self, sections: dict) -> list[str]:
        text = sections.get("certifications", "")
        certs = []
        for line in text.splitlines():
            line = re.sub(r"^[•·▪▸►\-\*]\s*", "", line.strip())
            if line and len(line) > 3:
                certs.append(line)
        return certs

    def _estimate_years_experience(self, experience: list[dict]) -> float:
        """Rough heuristic: count distinct years mentioned in experience."""
        year_re = re.compile(r"(20\d{2}|19\d{2})")
        all_years: set[int] = set()
        for entry in experience:
            for y in year_re.findall(entry.get("raw", "")):
                all_years.add(int(y))
        if len(all_years) >= 2:
            return float(max(all_years) - min(all_years))
        return float(len(experience))  # fallback: number of roles
